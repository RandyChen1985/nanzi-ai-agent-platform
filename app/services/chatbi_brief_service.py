"""Evidence-backed ChatBI business brief: AI report body + deterministic appendix."""

from __future__ import annotations

import logging
import re
import tempfile
from pathlib import Path
from typing import Any, Iterable, Mapping

from docx import Document

logger = logging.getLogger(__name__)

_ROW_KEYS = ("rows", "data", "result", "results", "items", "records")
_CHART_BLOCK_RE = re.compile(r"```chart[\s\S]*?```", re.IGNORECASE)
_FENCE_RE = re.compile(r"^```[\w-]*\s*$", re.MULTILINE)


class UnsupportedBriefClaim(ValueError):
    pass


class BriefInputError(ValueError):
    pass


def _rows(payload: Any, depth: int = 0) -> list[dict[str, Any]]:
    if depth > 4:
        return []
    if isinstance(payload, list):
        return [row for row in payload if isinstance(row, dict)]
    if not isinstance(payload, dict):
        return []
    for key in _ROW_KEYS:
        if key not in payload:
            continue
        value = payload[key]
        if isinstance(value, list):
            dict_rows = [row for row in value if isinstance(row, dict)]
            if dict_rows:
                return dict_rows
            columns = payload.get("columns")
            if isinstance(columns, list):
                return [
                    {
                        str(column.get("name") if isinstance(column, dict) else column): row[index]
                        for index, column in enumerate(columns)
                        if index < len(row)
                    }
                    for row in value
                    if isinstance(row, list)
                ]
            if value == []:
                return []
        elif isinstance(value, (dict, list)):
            nested = _rows(value, depth + 1)
            if nested:
                return nested
    return []


def _numeric_summary(rows: list[dict[str, Any]]) -> dict[str, dict[str, float]]:
    columns = {
        key
        for row in rows
        for key, value in row.items()
        if isinstance(value, (int, float)) and not isinstance(value, bool)
    }
    summary: dict[str, dict[str, float]] = {}
    for column in sorted(columns):
        values = [float(row[column]) for row in rows if isinstance(row.get(column), (int, float))]
        if values:
            summary[column] = {
                "sum": round(sum(values), 4),
                "average": round(sum(values) / len(values), 4),
                "min": min(values),
                "max": max(values),
            }
    return summary


def _numbers_in_text(text: str) -> set[float]:
    found: set[float] = set()
    for item in re.findall(r"-?\d+(?:\.\d+)?", str(text or "")):
        try:
            found.add(float(item))
        except ValueError:
            continue
    return found


def _supported_numbers(result_ref: Mapping[str, Any], assistant_report: str = "") -> set[float]:
    rows = _rows(result_ref.get("rows"))
    supported = {
        float(value)
        for row in rows
        for value in row.values()
        if isinstance(value, (int, float)) and not isinstance(value, bool)
    }
    for stats in _numeric_summary(rows).values():
        supported.update(float(value) for value in stats.values())
    supported.update(_numbers_in_text(assistant_report))
    return supported


def _validate_claims(claims: Iterable[str], supported_numbers: set[float]) -> None:
    for claim in claims:
        numbers = [float(item) for item in re.findall(r"-?\d+(?:\.\d+)?", str(claim))]
        if numbers and supported_numbers and any(number not in supported_numbers for number in numbers):
            raise UnsupportedBriefClaim(f"简报主张缺少数据证据：{claim}")


def sanitize_assistant_report(text: str, *, max_chars: int = 48_000) -> str:
    cleaned = str(text or "").strip()
    if not cleaned:
        return ""
    cleaned = _CHART_BLOCK_RE.sub("\n\n> （交互图表见在线会话记录，导出简报不含图文件。）\n\n", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    if len(cleaned) > max_chars:
        cleaned = cleaned[:max_chars] + "\n\n... [分析正文过长已截断]"
    return cleaned


def _brief_title(result_ref: Mapping[str, Any]) -> str:
    question = str(result_ref.get("question") or "数据分析").strip()
    title_core = re.sub(r"^(请|帮我|查询|查一下|查看)", "", question).strip() or "数据分析"
    return f"{title_core}业务简报"


def _build_evidence_appendix_lines(result_ref: Mapping[str, Any]) -> list[str]:
    rows = _rows(result_ref.get("rows"))
    numeric = _numeric_summary(rows)
    question = str(result_ref.get("question") or "数据分析").strip()
    lines = [
        "## 数据说明与证据附录",
        "",
        f"- 原始问题：{question}",
        f"- 数据集：{result_ref.get('dataset_name') or '未标注'}",
        f"- 数据行数：{len(rows)}",
        f"- 证据引用：ChatBI 结果 `{result_ref.get('result_id') or 'legacy'}`",
    ]
    sql = str(result_ref.get("sql") or "").strip()
    if sql:
        preview = sql if len(sql) <= 500 else sql[:500] + "..."
        lines.extend(["", "### SQL 摘要", "", f"```sql\n{preview}\n```"])
    if numeric:
        lines.extend(["", "### 结构化汇总（可核对）", ""])
        for column, stats in numeric.items():
            lines.append(
                f"- {column}：合计 {stats['sum']:g}，平均 {stats['average']:g}，"
                f"最低 {stats['min']:g}，最高 {stats['max']:g}"
            )
    elif rows:
        lines.extend(["", "- 本次结果含明细行，但未识别到可汇总的数值字段。"])
    else:
        lines.extend(["", "- 本次未附带结构化行数据，正文以上一轮 AI 分析为准。"])
    return lines


def _build_rows_only_markdown(result_ref: Mapping[str, Any], title: str) -> list[str]:
    rows = _rows(result_ref.get("rows"))
    numeric = _numeric_summary(rows)
    question = str(result_ref.get("question") or "数据分析").strip()
    lines = [
        f"# {title}",
        "",
        "## 分析范围",
        f"- 原始问题：{question}",
        f"- 数据集：{result_ref.get('dataset_name') or '未标注'}",
        f"- 数据行数：{len(rows)}",
        "",
        "## 核心数据",
    ]
    if numeric:
        for column, stats in numeric.items():
            lines.append(
                f"- {column}：合计 {stats['sum']:g}，平均 {stats['average']:g}，"
                f"最低 {stats['min']:g}，最高 {stats['max']:g}"
            )
    else:
        lines.append("- 本次结果不包含可汇总的数值字段，请以明细表为准。")
    summary = result_ref.get("result_summary")
    summary_text = summary if isinstance(summary, str) and summary.strip() else ""
    lines.extend([
        "",
        "## 业务结论",
        f"- {summary_text or '已完成本次数据查询，结论仅基于上述结构化结果。'}",
        "",
        "## 后续动作",
        "- 复核异常值及业务口径。",
        "- 对关键指标设置周期性监控和阈值告警。",
    ])
    return lines


def compose_business_brief_markdown(
    result_ref: Mapping[str, Any],
    *,
    assistant_report: str | None = None,
) -> str:
    title = _brief_title(result_ref)
    appendix = _build_evidence_appendix_lines(result_ref)
    report = sanitize_assistant_report(assistant_report or "")
    if report:
        body = report
        if not body.lstrip().startswith("#"):
            body = f"# {title}\n\n{body}"
        return "\n".join([body, "", "---", ""] + appendix)
    return "\n".join(_build_rows_only_markdown(result_ref, title) + ["", "---", ""] + appendix)


def build_business_brief(
    result_ref: Mapping[str, Any],
    *,
    assistant_report: str | None = None,
    requested_claims: Iterable[str] = (),
) -> dict[str, Any]:
    rows = _rows(result_ref.get("rows"))
    report = sanitize_assistant_report(assistant_report or "")
    if not rows and not report:
        raise BriefInputError(
            "缺少可导出的分析内容与结构化数据。请在本轮查数成功的回复上使用「生成业务简报」。"
        )
    supported = _supported_numbers(result_ref, report)
    _validate_claims(requested_claims, supported)
    title = _brief_title(result_ref)
    numeric = _numeric_summary(rows)
    context = dict(result_ref.get("analysis_context") or {})
    facts = {
        "row_count": len(rows),
        "numeric_summary": numeric,
        "analysis_context": context,
        "has_assistant_report": bool(report),
    }
    markdown = compose_business_brief_markdown(result_ref, assistant_report=report or None)
    return {
        "version": 2,
        "title": title,
        "facts": facts,
        "evidence": [{"result_id": result_ref.get("result_id"), "dataset_name": result_ref.get("dataset_name")}],
        "markdown": markdown,
    }


async def polish_business_brief_markdown(
    markdown: str,
    *,
    supported_numbers: set[float],
) -> str:
    """Optional LLM pass to tidy layout without inventing numbers; falls back on any error."""
    from app.services.ai.config import AgentConfigProvider
    from app.services.ai.runtime.agentscope.chat import chat_client_from_handle
    from app.services.ai.runtime.agentscope.messages import system_user_prompt_messages

    try:
        llm = await AgentConfigProvider.get_configured_llm(streaming=False)
        chat_client = chat_client_from_handle(llm)
        content = await chat_client.generate_text(
            system_user_prompt_messages(
                (
                    "你是企业数据分析简报编辑。请把用户提供的 Markdown 整理为结构清晰的业务简报。"
                    "规则：不得新增任何数字或事实；不得删除表格中的关键行；可合并重复段落、统一标题层级；"
                    "保留原有结论与建议；图表占位说明可保留。"
                    "只输出 Markdown 正文，不要用代码围栏包裹全文。"
                ),
                user_prompt=markdown,
            )
        )
        polished = str(content or "").strip()
        polished = re.sub(r"^```(?:markdown)?\s*|\s*```$", "", polished, flags=re.IGNORECASE).strip()
        if not polished:
            return markdown
        _validate_claims([polished], supported_numbers)
        return polished
    except UnsupportedBriefClaim:
        logger.info("[ChatBIBrief] LLM polish rejected: unsupported numbers, using draft")
        return markdown
    except Exception as exc:
        logger.info("[ChatBIBrief] LLM polish skipped: %s", exc)
        return markdown


async def build_business_brief_async(
    result_ref: Mapping[str, Any],
    *,
    assistant_report: str | None = None,
    polish: bool = True,
) -> dict[str, Any]:
    brief = build_business_brief(result_ref, assistant_report=assistant_report)
    if not polish:
        return brief
    supported = _supported_numbers(result_ref, sanitize_assistant_report(assistant_report or ""))
    brief["markdown"] = await polish_business_brief_markdown(brief["markdown"], supported_numbers=supported)
    return brief


def publish_business_brief_docx(brief: Mapping[str, Any]):
    """Create a private 24-hour Word artifact from brief Markdown."""
    from app.services.ai.tools.generated_file_service import publish

    safe_title = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", str(brief.get("title") or "业务简报"))[:60]
    with tempfile.TemporaryDirectory(prefix="chatbi_brief_") as temp_dir:
        path = Path(temp_dir) / f"{safe_title}.docx"
        document = Document()
        for line in str(brief.get("markdown") or "").splitlines():
            if line.startswith("# "):
                document.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                document.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                document.add_heading(line[4:], level=3)
            elif line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("> "):
                document.add_paragraph(line[2:])
            elif line.strip() == "---":
                document.add_paragraph("")
            elif _FENCE_RE.match(line):
                continue
            elif line.strip():
                document.add_paragraph(line)
        document.save(path)
        return publish(path, path.name)
